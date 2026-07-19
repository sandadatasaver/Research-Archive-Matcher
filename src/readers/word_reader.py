import docx
import os
import logging

logger = logging.getLogger(__name__)

class WordReader:
    """
    A class to read Word document (.docx) files, extracting paragraphs and tables
    to read lists of target publications.
    """
    def __init__(self, file_path):
        self.file_path = file_path
        if not os.path.exists(file_path):
            raise FileNotFoundError(f"Word file not found: {file_path}")

    def read_lines(self) -> list:
        """
        Extracts all paragraph text and table cell text as a list of non-empty strings.
        """
        lines = []
        try:
            doc = docx.Document(self.file_path)
            # Read paragraphs
            for para in doc.paragraphs:
                text = para.text.strip()
                if text:
                    lines.append(text)
            
            # Read tables
            for table in doc.tables:
                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        val = cell.text.strip()
                        if val and val not in row_text:
                            row_text.append(val)
                    if row_text:
                        lines.append(" | ".join(row_text))
        except Exception as e:
            logger.error(f"Error reading Word file {self.file_path}: {e}")
        return lines
