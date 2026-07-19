import docx
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
import os
import datetime

class WordReporter:
    """
    Generates a beautifully formatted Microsoft Word (.docx) matching summary report.
    """
    @staticmethod
    def generate_report(results: dict, db_stats: dict, output_path: str = "matching_report.docx") -> str:
        doc = docx.Document()
        
        # Set margins
        for section in doc.sections:
            section.top_margin = Inches(1)
            section.bottom_margin = Inches(1)
            section.left_margin = Inches(1)
            section.right_margin = Inches(1)
            
        # Add Header Style / Colors
        style_title = doc.styles.add_style('RAM_Title', docx.enum.style.WD_STYLE_TYPE.PARAGRAPH)
        style_title.font.name = 'Calibri'
        style_title.font.size = Pt(24)
        style_title.font.bold = True
        style_title.font.color.rgb = RGBColor(31, 78, 121) # Classic Navy
        
        style_h1 = doc.styles.add_style('RAM_H1', docx.enum.style.WD_STYLE_TYPE.PARAGRAPH)
        style_h1.font.name = 'Calibri'
        style_h1.font.size = Pt(16)
        style_h1.font.bold = True
        style_h1.font.color.rgb = RGBColor(31, 78, 121)
        
        # Document Title
        p = doc.add_paragraph('Research Archive Matcher (RAM)', style='RAM_Title')
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        p_sub = doc.add_paragraph(f'Matching Execution Report — Generated on {datetime.date.today().strftime("%B %d, %Y")}')
        p_sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_sub.runs[0].font.italic = True
        p_sub.runs[0].font.size = Pt(11)
        
        doc.add_paragraph().paragraph_format.space_after = Pt(20)
        
        # Section 1: Executive Summary
        doc.add_paragraph('1. Executive Summary', style='RAM_H1')
        
        matched_count = len(results.get("matched", []))
        unmatched_count = len(results.get("unmatched", []))
        total_targets = matched_count + unmatched_count
        match_rate = (matched_count / total_targets * 100) if total_targets > 0 else 0
        
        p_sum = doc.add_paragraph()
        p_sum.add_run('The matching process was executed successfully. Below is a high-level overview of the library index and matched publications:\n\n').font.size = Pt(11)
        
        # Add bullet points
        bullet_style = 'List Bullet'
        doc.add_paragraph(f'Total PDFs in Local Library Index: {db_stats.get("total_docs", 0)}', style=bullet_style)
        doc.add_paragraph(f'Total Target Publications to Match: {total_targets}', style=bullet_style)
        doc.add_paragraph(f'Successful Matches: {matched_count}', style=bullet_style)
        doc.add_paragraph(f'Unmatched Publications: {unmatched_count}', style=bullet_style)
        doc.add_paragraph(f'Success Match Rate: {match_rate:.1f}%', style=bullet_style)
        
        doc.add_paragraph().paragraph_format.space_after = Pt(15)
        
        # Section 2: Detailed Matches Table
        doc.add_paragraph('2. Successful Publication Matches', style='RAM_H1')
        doc.add_paragraph('The following table details the target citations that were successfully matched against documents in your local PDF library.')
        
        matched_list = results.get("matched", [])
        if matched_list:
            table = doc.add_table(rows=1, cols=4)
            table.alignment = WD_TABLE_ALIGNMENT.CENTER
            table.style = 'Light Shading Accent 1'
            
            hdr_cells = table.rows[0].cells
            hdr_cells[0].text = 'Target Citation'
            hdr_cells[1].text = 'Matched PDF Title'
            hdr_cells[2].text = 'Score'
            hdr_cells[3].text = 'Method'
            
            for item in matched_list:
                row_cells = table.add_row().cells
                row_cells[0].text = item["target_title"][:60] + ("..." if len(item["target_title"]) > 60 else "")
                row_cells[1].text = item["matched_title"][:60] + ("..." if len(item["matched_title"]) > 60 else "")
                row_cells[2].text = f'{item["score"]}%'
                row_cells[3].text = item["method"]
                
                # Make scores center-aligned
                row_cells[2].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                row_cells[3].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        else:
            doc.add_paragraph('No matches were found above the confidence threshold. All publications remain unmatched.', style='Normal')
            
        doc.add_paragraph().paragraph_format.space_after = Pt(15)
        
        # Section 3: Unmatched Publications
        doc.add_paragraph('3. Unmatched Publications', style='RAM_H1')
        doc.add_paragraph('The following publications could not be matched against any document in your local PDF archive above the similarity threshold.')
        
        unmatched_list = results.get("unmatched", [])
        if unmatched_list:
            table_un = doc.add_table(rows=1, cols=3)
            table_un.alignment = WD_TABLE_ALIGNMENT.CENTER
            table_un.style = 'Light Shading Accent 1'
            
            hdr_un = table_un.rows[0].cells
            hdr_un[0].text = 'Unmatched Publication'
            hdr_un[1].text = 'Best Index Candidate'
            hdr_un[2].text = 'Best Score'
            
            for item in unmatched_list:
                row_cells = table_un.add_row().cells
                row_cells[0].text = item["target_title"][:80] + ("..." if len(item["target_title"]) > 80 else "")
                row_cells[1].text = item["best_candidate_title"][:80] + ("..." if len(item["best_candidate_title"]) > 80 else "")
                row_cells[2].text = f'{item["best_candidate_score"]}%'
                row_cells[2].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        else:
            doc.add_paragraph('Excellent! Every single target publication was matched with full confidence.', style='Normal')
            
        # Save Document
        doc.save(output_path)
        return output_path
